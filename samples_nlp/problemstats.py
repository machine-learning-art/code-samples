import os, sys, csv, re, glob
from docx import *
from theorystats import *
from bs4 import BeautifulSoup
import zipfile



#Obtained from http://etienned.github.io/posts/extract-text-from-word-docx-simply/
#It extracts all the text in a docx file (including text in tables)
try:
    from xml.etree.cElementTree import XML
except ImportError:
    from xml.etree.ElementTree import XML
import zipfile
"""
Module that extract text from MS XML Word document (.docx).
(Inspired by python-docx <https://github.com/mikemaccana/python-docx>)
"""
WORD_NAMESPACE = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
PARA = WORD_NAMESPACE + 'p'
TEXT = WORD_NAMESPACE + 't'
 
def get_docx_text(path):
    """
    Take the path of a docx file as argument, return the text in unicode.
    """
    document = zipfile.ZipFile(path)
    xml_content = document.read('word/document.xml')
    document.close()
    tree = XML(xml_content)
 
    paragraphs = []
    for paragraph in tree.getiterator(PARA):
        texts = [node.text
                 for node in paragraph.getiterator(TEXT)
                 if node.text]
        if texts:
            paragraphs.append(''.join(texts))
 
    return '\n\n'.join(paragraphs)

	
	
#input:'docpath' is the document's path
#output: returns the number of interactive steps, number of hints, number of second layers in an array: [IScount, Hintcount, secondlayerc]
def getSolstats(docpath):
    doc = Document(docpath)
    IScount = 0
    Hintcount = 0
    secondlayerc = 0
    for par in doc.paragraphs:
        if "[IS" in par.text:
            IScount += 1
        if "Hint" in par.text:
            Hintcount +=1
        if '[?]' in par.text or '(?)' in par.text:
            secondlayerc += 1
    return [IScount, Hintcount, secondlayerc]
#getSolstats('C:\\Users\\eabalo\\Desktop\\TimingSTAAR\\Code\\Sample problem scripts\\TEKS303E-pr001.docx')


#Extracting problem statement text by extracting text in the first 1x1 table in the script
def getProbtext1(docpath):
    doc = Document(docpath)
    text = ''
    for tab in doc.tables:
        #if len(tab.rows)==1 and len(tab.columns)==1: #1x1 table to exclude other tables
        if len(tab.columns)==1: #1x1 table to exclude other tables
            for row in tab.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        text += ' ' + par.text
            break #stop after finding the first 1x1 table
        else:
            pass
    text = text.replace(u'\u2019',"'")
    text = text.encode('ascii','ignore')
    if 'Hint' in text:
        text = "ERROR! Hint in the text."
    elif 'Answer' in text:
        text = "ERROR! Answer in the text."
    elif 'DS' in text:
        text = "ERROR! DS in the text."
    return text
#getProbtext1('C:\\Users\\eabalo\\Desktop\\TimingSTAAR\\Code\\Sample problem scripts\\TEKS302B-pr001.docx')


#Extracting problem statement text by extracting text between keywords 'Problem' and the section keyword that follows immediately ('Dataset' or 'Hint' etc.)
def getProbtext2(docpath):
    doc = Document(docpath)
    problemstruct = []
    parnum = 0 # paragraph number
    text = ''
    for par in doc.paragraphs:
        if re.search('^Problem', par.text):
            problemstruct .append(['p', parnum])
        if re.search('^(Dataset|Datasets)', par.text):
            problemstruct .append(['d', parnum])
        if re.search('^Hint', par.text):
           problemstruct .append(['h', parnum])
        if re.search('^Solution', par.text):
            problemstruct .append(['s', parnum])
        if re.search('^Answer', par.text):
            problemstruct .append(['a', parnum])
        parnum += 1
    problemstruct.sort(key=lambda e: e[1]) #sorting on the paragraph number
    probpar = problemstruct[0][1] #paragraph number for "Problem"
    nextpar = problemstruct[1][1] #paragraph number for next section following "Problem"
    for par in doc.paragraphs[probpar+1:nextpar]:
        text += ' ' + par.text
    text = text.replace(u'\u2019',"'")
    text = text.encode('ascii','ignore')
    if 'Hint' in text:
        text = "ERROR! Hint in the text."
    elif 'Answer' in text:
        text = "ERROR! Answer in the text."
    elif 'DS' in text:
        text = "ERROR! DS in the text."
    return text
#getProbtext2('C:\\Users\\eabalo\\Desktop\\TimingSTAAR\\Code\\Sample problem scripts\\TEKS302B-pr001.docx')


#extracts the problem statement using get_docx_text 
#text is extracted between keywords 'Problem' and 'Dataset'
def getProbtext3(docpath):
    text = get_docx_text(docpath)
    text = text.replace(u'\u2019',"'")
    text = text.encode('ascii','ignore')
    try:
        list1 = [m.start(0) for m in re.finditer('Problem', text)] #indices of 'Problem'
        list2 = [m.start(0) for m in re.finditer('Dataset|Datasets', text)] #indices of 'Dataset'
        list3 = [n for n in list1 if n < list2[-1]] #'Problem' index just before last 'Dataset'
        ind2 = text.index('\n', list3[-1]+1) #index of newline character after last 'Problem'
        output = text[ind2:list2[-1]]
    except:
        output = 'ERROR!'
    if 'Hint' in output:
        output = "ERROR! Hint in the text."
    elif 'Answer' in output:
        output = "ERROR! Answer in the text."
    elif 'DS' in output:
        output = "ERROR! DS in the text."
    return output
#getProbtext3('C:\\Users\\eabalo\\Desktop\\TimingSTAAR\\Code\\Sample problem scripts\\TEKS303E-pr001.docx')


#input:
#output:
#subroutines: getProbtext1() and getProbtext2()
def getProbtext(docpath):
    try:
        text = getProbtext1(docpath) #look for 1x1 table
        if len(text)==0 or 'Answer' in text or 'Hint' in text or 'DS' in text or 'Dataset' in text \
        or 'Reuse' in text or 'ERROR' in text:
            text = getProbtext2(docpath)
            if len(text)==0 or 'Answer' in text or 'Hint' in text or 'DS' in text or 'Dataset' in text \
            or 'Reuse' in text or 'ERROR' in text:
                text = getProbtext3(docpath)
                if len(text)==0 or 'Answer' in text or 'Hint' in text or 'DS' in text or 'Dataset' in text or 'ERROR' in text:
                    text = 'ERROR!'
                
    except:
        text = 'ERROR!'
        
    return text
#getProbtext('C:\\Users\\eabalo\\Desktop\\TimingSTAAR\\Code\\Sample problem scripts\\TEKS302B-pr001.docx')


#removing animator notes in script
def removeBracketed(text):
    '''Remove text enclosed in square brackets.  Regexes can't really handle
    nested brackets, so this does it manually.'''
    bc = 0
    temp = ''
    for char in text:
        if char == '[':
            bc += 1
        if bc == 0: temp += char
        if char == ']':
            bc -= 1
    return temp
#print removeBracketed(\getProbtext('C:\\Users\\eabalo\\Desktop\\TimingSTAAR\\Code\\Sample problem scripts\\TEKS302B-pr001.docx')).split()


def processText(text):
    #text = re.sub('Problem', '', text)
    #text = re.sub('Datasets', '', text)
    #text = re.sub('Solution', '', text)
    #text = re.sub('Hint \d?', '', text)
    try:
        text = text.replace(u'\u2019',"'")
        text = text.encode('ascii','ignore')
        text = re.sub('([A-Z]|[a-z])(\)|\:|\.)', '', text) #removing answer choice numbering labels for the form A), a.
        text = re.sub('\;|\,|\.|\:|\?|\!', '', text)
        text = removeBracketed(text)
        text = text.replace('  ', ' ')
        #text = re.sub('\[\?\]', '', text) #removing second layer animator notes
        text = re.sub('\(\?\)', '', text)
        text = re.sub('^\n', '', text)
    except:
        text = "ERROR!"
    if 'DS' in text:
        text = "ERROR! DS in text"
    return text


#returns the number of words in the problem statement
def probCount(docpath):
    #doc = Document(docpath)
    text = getProbtext(docpath)
    output = ''
    if text and 'ERROR' in text:
        output = -1
    else:
        output = len(processText(text).split())
    return output
	
	
#returns the solution text
def getSolution(docpath):
	doc = Document(docpath)
	text = ''
	tblen = len(doc.tables) #number of tables in the script
	for par in doc.paragraphs:  
		style = par.style #style applied to text
		for run in par.runs:
			if not run.strike:
			#if not run.font.strike: #in a different version of python-docx
				for t in run._r.t_lst:
					if t.text is None:
						pass
					else:
						text += ' ' + t.text
        					
	if len(doc.tables)>3: #text in the second layers
		for j in range(3, tblen):
			table = doc.tables[j]
			for row in table.rows:
				for cell in row.cells:
					for par in cell.paragraphs:
						style = par.style
						for run in par.runs:
							if not run.strike:
							#if not run.font.strike: #in a different version of python-docx
								for t in run._r.t_lst:
									if t.text is None:
										pass
									else:
										text += ' ' + t.text
	#text = re.sub('^ ', ' ', text)
	# if 'Hint' in text:
		# text = "ERROR! Hint in the text."
	# elif 'Problem' in text:
		# text = "ERROR! Problem in the text."
	text = text.replace(u'\u2019',"'")
	text = text.encode('ascii','ignore')
	try:
		output = text.split('Solution')[1] #split the text after 'Solution'
	except:
		output = "ERROR!"
	else:
		return output
#print solCount('C:\\Users\\eabalo\\Desktop\\TimingSTAAR\\Code\\Sample problem scripts\\TEKS302A-pr001.docx')
#print processText(getSolution('C:\\Users\\eabalo\\Desktop\\TimingSTAAR\\Code\\Sample problem scripts\\TEKS302A-pr001.docx'))


#returns the number of words in the solution
def solCount(docpath):
    #doc = Document(docpath)
    text = getSolution(docpath)
    output = ''
    if text and 'ERROR' in text:
        output = -1
    else:
        output = len(processText(text).split())
    return output
#print solCount('C:\\Users\\eabalo\\Desktop\\TimingSTAAR\\Code\\Sample problem scripts\\TEKS302A-pr001.docx')
#print processText(getSolution2('C:\\Users\\eabalo\\Desktop\\TimingSTAAR\\Code\\Sample problem scripts\\TEKS302A-pr001.docx'))


#returns the number of second layers
#It checks for paragraph borders, [?], and tables after 'Solution' in the given order and works with the syntax detected first
#CURRENTLY NOT USED
def seclayerCount(docpath):
    document = zipfile.ZipFile(docpath)
    xml_content = document.read('word/document.xml')
    document.close()
    soup = BeautifulSoup(xml_content, 'xml')
    
    numseclayers = 0
    
    #find tag containing 'Solution' at the beginning of the line
    solutionheader = soup.find(text = re.compile('^Solution')) 
    pBdrtags = solutionheader.find_all_next('pBdr')
    qtags = solutionheader.find_all_next(text='[?]')
    tbltags = solutionheader.find_all_next('tbl')
    
    if pBdrtags: #if boxes are used around second layer text
        for par in pBdrtags: #look for 'pBdr' tag
            if par.find_parent('p').find('t'): #go up to parent tag 'p' then down to 't' tag
                numseclayers += 1
    elif qtags: #if [?] are used in the script
        for par in qtags:
            numseclayers += 1
    elif tbltags: #if tables are used
        for par in tbltags:
            if par.find('t'):
                numseclayers += 1
    return numseclayers
	
#print seclayerCount('C:\\Users\\eabalo\\Documents\\TimingSTAAR\\Code\\Sample problem scripts\\TEKS302A-pr001.docx')
#print seclayerCount('C:\\Users\\eabalo\\Documents\\TimingSTAAR\\Code\\Sample problem scripts\\TEKS604A-150pr.docx')