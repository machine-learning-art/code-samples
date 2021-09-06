import os, sys, csv, re, glob
from docx import *


#input: docpath is the document's path
#output: returns all the text except the header info 'TEKS302A-th001'
def getText(docpath):
    doc = Document(docpath)
    text = ''
    for par in doc.paragraphs[1:]:
        style = par.style
        #text += '' + par.text
        for run in par.runs:
            if not run.strike:
			#if not run.font.strike: #in a different version of python-docx?
                for t in run._r.t_lst:
                    if t.text is None:
                        pass
                    else:
                        text += '' + t.text    #why add white spaces?
        text = re.sub('^ ','',text)
    return text
#print getText('Sample lesson scripts\\TEKS402A-th001.docx')


#input: 'docpath' is the document's path
#output: returns number of 'Screen #', [Next], and [Submit] as elements of an array
def getStats(docpath):
    doc = Document(docpath)
    branchcount = 0 #number of Correct and Incorrect branches
    corrcount = 0 #number of Correct branches
    screencount = []
    nextcount = []
    submitcount = []
    for par in doc.paragraphs[1:]:
        screencount += re.findall('^Screen', par.text)
        nextcount += re.findall('\[next', par.text.lower())
        submitcount += re.findall('\[submit', par.text.lower())
        #branchcout += re.findall('^(correct|incorrect)', par.text.lower())
        #corrcount += re.findall('^correct', par.text.lower())
        if re.match('^correct',par.text.lower()) or re.match('^incorrect', par.text.lower()):
            branchcount += 1
            if re.match('^correct',par.text.lower()):
                corrcount += 1
    return [len(screencount), len(nextcount), len(submitcount), branchcount, corrcount]
#print getStats('Sample lesson scripts\\TEKS402A-th001.docx')    


#input: 'docpath' is the document's path
#output: returns the number of tables in the document (NB: Some tables do not have data)
def getTablecount(docpath):
    doc = Document(docpath)
    return len(doc.tables)
#print getTablecount('Sample lesson scripts\\TEKS402A-th001.docx')
 
 
#input: text
#output: returns a 'cleaned' version of the text by removing text in brackets
def removeBracketed(text):
    '''Remove text enclosed in square brackets.  Regexes can't really handle
    nested brackets, so this does it manually.'''
    bc = 0
    temp = ''
    for char in text:
        if char == '[':
            bc += 1
        if bc == 0: temp += char
        if char == ']':
            bc -= 1
    return temp

	
#input: 'docpath' is the document's path
#output: returns on-screen text
#subroutines: getText()
def onscreenText(docpath):
    text = getText(docpath)
    try:
        text = text.replace(u'\u2019',"'")
        text = text.encode('ascii','ignore')
        #text = re.sub('([A-Z]|[a-z])(\)|\.|\:)', '', text) #removing answer choice numbering labels for the form A), a.
        text = re.sub('L\d*\.\d*', '', text) #remove line labels L1.1
        text = removeBracketed(text)
        text = text.replace('  ', ' ')
        text = re.sub('\.|\,|\;|\?|\!|\:', '', text)
        #text = re.sub('\[\?\]', '', text) #removing second layer animator notes
        #text = re.sub('\(\?\)', '', text)
        text = re.sub('^\n', '', text)
        text = re.sub('^Screen \d', '', text)
        text = re.sub('Correct|Incorrect', '', text)
    except:
        text = "ERROR!"
    return text
#print onscreenText('Sample lesson scripts\\TEKS402A-th001.docx')

#input: 'docpath' is the document's path
#output: returns the total number of words in all tables
def getTablewordcount(docpath):
    output = 0
    text = ''
    try:
        doc = Document(docpath)
        text = ''
        for tab in doc.tables:
            for row in tab.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        text += ' ' + par.text
        output = len(text.split())
    except:
        output = "ERROR!"
    return output
#print getTablewordcount('Sample lesson scripts\\TEKS402A-th001.docx')