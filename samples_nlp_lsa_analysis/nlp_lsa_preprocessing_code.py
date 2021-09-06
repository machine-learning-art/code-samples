from docx import *
import re, glob, os


#CODE TO EXTRACT TEXT FROM DOCX --------------------------------------------------
#returns all the text except the header info 'TEKS302A-th001'
def getText(docpath):
    text = ''
    try:
        doc = Document(docpath)
        for par in doc.paragraphs[1:]:
            style = par.style
            #text += '' + par.text
            for run in par.runs:
                if not run.strike:
                    for t in run._r.t_lst:
                        if t.text is None:
                            pass
                        else:
                            text += '' + t.text    #why add white spaces?
            text = text.replace(u'\u2019',"'")
            text = text.encode('ascii','ignore')
            text = re.sub('^ ','',text)
            text = text + ' ' #insert white space after each paragraph to avoid sticking words together
            #text = text.replace(' ', ' ')
    except:
        pass
    return text
#print getText('Sample lesson scripts\\TEKS402A-th001.docx')


#Extracting problem statement text by extracting text in the first 1x1 table in the script
#extracts text outside of tables by calling getText
def getText4(docpath):
    doc = Document(docpath)
    text = ''
    for tab in doc.tables:
        #if len(tab.rows)==1 and len(tab.columns)==1: #1x1 table to exclude other tables
        if len(tab.columns)==1: #1x1 table to exclude other tables
            for row in tab.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        for run in par.runs:
                            if not run.strike:
                                for t in run._r.t_lst:
                                    if t.text is None:
                                        pass
                                    else:
                                        text += ' ' + t.text    #why add white spaces?
                        #text += ' ' + par.text
                        #style = par.style
            break #stop after finding the first 1x1 table
            
    text += getText(docpath) #adding text outside tables
    
    text = text.replace(u'\u2019',"'")
    text = text.encode('ascii','ignore')
    text = re.sub('^ ','',text)
    return text

	
#CODE TO PROCESS TEXT ------------------------------------------------------
#removes comments in brackets
def removeBracketed(text):
    '''Remove text enclosed in square brackets.  Regexes can't really handle
    nested brackets, so this does it manually.'''
    bc = 0
    temp = ''
    for char in text:
        if char == '[':
            bc += 1
        if bc == 0: temp += char
        if char == ']':
            bc -= 1
    return temp

#removes all punctuations
def remove_punc(text):
    punc = '!@#$%^&*()_-+={}[]:;"\'|<>,.?/~`'
    for char in punc:
        text = text.replace(char, ' ')
        text = text.replace('  ', ' ')
    return text

#removes digits
def remove_digits(text):
    digits = '0123456789'
    for char in digits:
        text = text.replace(char, '')
        text = text.replace('  ', ' ')
    return text

#cleans text by lowering case, removing bracketed text, 
#removing punctuations and digits
def clean_text(text):
    text = text.lower()
    text = removeBracketed(text)
    text = remove_punc(text)
    text = remove_digits(text)
    return text

#CODE TO BUILD A CORPUS---------------------------------------------------------
#returns the paths of theory and problem scripts in a lesson folder
#lessonpath is the lesson's folder path
#for example: lessonpath = 'C:/Users/eabalo/Documents/STAAR2014/3g/308A/'

def scriptpaths(lessonpath):
    #paths of docx files needed to create corpus
    paths = []

    for i in glob.glob(lessonpath + 'Scripts/*.docx'): #append theory scripts
        if '~' in i: #ignore corrupted files
            pass
        else:
            paths.append(i)
        
    for i in glob.glob(lessonpath + 'Problems/*.docx'): #append problem scripts
        if '~' in i: #ignore corrupted files
            pass
        else:
            paths.append(i)
    return paths

	
#CODE TO CREATE A CORPUS------------------------------------------------------------------
#creates a folder with text files built from text extracted from 
#theory and problem scripts in a given lesson
#input: lesson path in SVN folder; for example 
    #lessonpath = 'C:/Users/eabalo/Documents/STAAR2014/3g/308B/'
#output: named files containing text extracted from corresponding docx file
#calls scriptpaths(), getText4(), clean_text()

def corpusText(lessonpath):
    #lesson's path in SVN folder
    #lessonpath = 'C:/Users/eabalo/Documents/STAAR2014/3g/308B/'

    #corpus path to store processed txt files
    corpuspath = 'C:/Users/eabalo/Desktop/STAAR35Analyses/data/corpus/'

    #lesson number
    lesson = lessonpath.split('/')[-2]

    #path for folder of corpus (txt files)
    newpath = corpuspath + lesson + '/'

    #create folder
    if not os.path.exists(newpath): 
        os.makedirs(newpath)
    
    #paths of docx files in SVN folder (theory and problem) 
    filepaths = scriptpaths(lessonpath)

    #creating corpus of text from lesson scripts
    #extracting text from docx files
    for i in filepaths:
        textfilename = i.split('\\')[-1].replace('docx', 'txt')
        text = getText4(i)
        text = clean_text(text)
        with open(newpath + textfilename, 'wb') as f:
            f.write(text)